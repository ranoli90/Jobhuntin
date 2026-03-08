"""Enhanced document processing: PDF, DOC/DOCX, and OCR support."""

import io
import os
from fastapi import HTTPException

import fitz  # PyMuPDF for PDF
import docx  # python-docx for DOCX
from PIL import Image
import pytesseract  # OCR support

from shared.logging_config import get_logger

logger = get_logger("sorce.document_processor")


class DocumentProcessor:
    """Enhanced document processor supporting PDF, DOC/DOCX, and OCR."""

    SUPPORTED_PDF_TYPES = {
        "application/pdf",
        "application/x-pdf",
    }

    SUPPORTED_DOCX_TYPES = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
        "application/vnd.ms-word",
    }

    SUPPORTED_IMAGE_TYPES = {
        "image/jpeg",
        "image/jpg",
        "image/png",
        "image/tiff",
        "image/bmp",
    }

    def __init__(self):
        """Initialize the document processor."""
        self.ocr_enabled = self._check_ocr_availability()
        if self.ocr_enabled:
            logger.info("OCR is available for scanned document processing")
        else:
            logger.warning(
                "OCR not available - pytesseract not installed or Tesseract not found"
            )

    def _check_ocr_availability(self) -> bool:
        """Check if OCR is available."""
        try:
            # Try to get Tesseract version
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False

    def get_file_type(self, filename: str, content_type: str) -> str:
        """Determine the document type from filename and content type."""
        # First check content type
        if content_type in self.SUPPORTED_PDF_TYPES:
            return "pdf"
        elif content_type in self.SUPPORTED_DOCX_TYPES:
            return "docx"
        elif content_type in self.SUPPORTED_IMAGE_TYPES:
            return "image"

        # Fallback to filename extension
        ext = os.path.splitext(filename.lower())[1]
        if ext == "pdf":
            return "pdf"
        elif ext in ["docx", "doc"]:
            return "docx"
        elif ext in ["jpg", "jpeg", "png", "tiff", "bmp"]:
            return "image"

        return "unknown"

    def is_supported_file(self, filename: str, content_type: str) -> bool:
        """Check if the file format is supported."""
        file_type = self.get_file_type(filename, content_type)
        return file_type in ["pdf", "docx", "image"]

    async def extract_text_from_document(
        self, file_bytes: bytes, filename: str, content_type: str, use_ocr: bool = True
    ) -> str:
        """Extract text from various document types with OCR fallback.

        Args:
            file_bytes: The file content as bytes
            filename: Original filename
            content_type: MIME content type
            use_ocr: Whether to attempt OCR for scanned documents

        Returns:
            Extracted text content

        Raises:
            HTTPException: If file format is not supported or processing fails
        """
        file_type = self.get_file_type(filename, content_type)

        if file_type == "pdf":
            return await self._extract_text_from_pdf(file_bytes, use_ocr)
        elif file_type == "docx":
            return await self._extract_text_from_docx(file_bytes)
        elif file_type == "image":
            return await self._extract_text_from_image(file_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_type}. Supported formats: PDF, DOCX, and images.",
            )

    async def _extract_text_from_pdf(
        self, pdf_bytes: bytes, use_ocr: bool = True
    ) -> str:
        """Extract text from PDF with OCR fallback for scanned documents."""
        try:
            # First attempt direct text extraction
            text_parts = []
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                if doc.page_count == 0:
                    raise HTTPException(
                        status_code=422, detail="The PDF file appears to be empty."
                    )

                # Check if PDF is likely scanned (has images but no text)
                is_scanned = self._is_pdf_scanned(doc)
                logger.info(
                    f"PDF detected as {'scanned' if is_scanned else 'text-based'}"
                )

                if is_scanned and use_ocr and self.ocr_enabled:
                    logger.info("Attempting OCR extraction from scanned PDF")
                    ocr_text = await self._ocr_extract_from_pdf(pdf_bytes)
                    if ocr_text.strip():
                        return ocr_text
                    # If OCR fails, fall back to direct extraction
                    logger.warning(
                        "OCR extraction failed, falling back to direct text extraction"
                    )

                # Direct text extraction
                for page in doc:
                    text_parts.append(page.get_text())

                extracted_text = "\n".join(text_parts)

                # If direct extraction yields very little text and OCR is enabled, try OCR
                if (
                    len(extracted_text.strip()) < 100
                    and use_ocr
                    and self.ocr_enabled
                    and not is_scanned
                ):
                    logger.info(
                        "Direct text extraction yielded minimal text, attempting OCR"
                    )
                    ocr_text = await self._ocr_extract_from_pdf(pdf_bytes)
                    if ocr_text.strip():
                        return ocr_text

                return extracted_text

        except fitz.FileDataError as e:
            logger.error(f"Failed to open PDF: {e}")
            raise HTTPException(
                status_code=422,
                detail="Invalid PDF format or corrupted file. Please ensure you're uploading a valid PDF.",
            )
        except Exception as e:
            logger.error(f"Unexpected error during PDF extraction: {e}")
            raise HTTPException(
                status_code=422,
                detail="Failed to process PDF. Please try a different version of your resume.",
            )

    def _is_pdf_scanned(self, doc) -> bool:
        """Check if PDF is likely scanned (has images but minimal text)."""
        try:
            # Check first few pages for text content
            text_content = ""
            for i, page in enumerate(doc):
                if i >= 3:  # Check only first 3 pages
                    break
                text_content += page.get_text()

            # If very little text but has images, likely scanned
            has_images = any(page.get_images() for page in doc)
            is_scanned = len(text_content.strip()) < 50 and has_images

            if is_scanned:
                logger.debug(
                    f"PDF appears scanned: {len(text_content.strip())} chars, {sum(len(page.get_images()) for page in doc)} images"
                )

            return is_scanned
        except Exception:
            return False

    async def _ocr_extract_from_pdf(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF using OCR."""
        try:
            # Convert PDF to images for OCR
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                all_text = []

                for page_num in range(len(doc)):
                    try:
                        # Get page as image
                        pix = page.get_pixmap()
                        img = Image.frombytes("RGB", pix.size, pix.samples)

                        # Perform OCR on the image
                        text = pytesseract.image_to_string(img)
                        if text.strip():
                            all_text.append(f"--- Page {page_num + 1} ---\n{text}")

                    except Exception as e:
                        logger.warning(f"Failed to OCR page {page_num + 1}: {e}")
                        continue

                return "\n".join(all_text)

        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""

    async def _extract_text_from_docx(self, docx_bytes: bytes) -> str:
        """Extract text from DOCX document."""
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))

            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise HTTPException(
                status_code=422,
                detail="Failed to process DOCX file. Please ensure the file is not corrupted.",
            )

    async def _extract_text_from_image(self, image_bytes: bytes) -> str:
        """Extract text from image using OCR."""
        if not self.ocr_enabled:
            raise HTTPException(
                status_code=422,
                detail="OCR is not available. Please install Tesseract to process image files.",
            )

        try:
            # Convert bytes to PIL Image
            image = Image.open(io.BytesIO(image_bytes))

            # Convert to RGB mode if necessary
            if image.mode != "RGB":
                image = image.convert("RGB")

            # Perform OCR
            text = pytesseract.image_to_string(image)
            return text

        except Exception as e:
            logger.error(f"Failed to extract text from image: {e}")
            raise HTTPException(
                status_code=422,
                detail="Failed to process image file. Please ensure the image is clear and readable.",
            )

    async def extract_metadata(
        self, file_bytes: bytes, filename: str, content_type: str
    ) -> dict:
        """Extract metadata from the document."""
        file_type = self.get_file_type(filename, content_type)
        metadata = {
            "file_type": file_type,
            "filename": filename,
            "file_size": len(file_bytes),
            "content_type": content_type,
        }

        if file_type == "pdf":
            metadata.update(await self._extract_pdf_metadata(file_bytes))
        elif file_type == "docx":
            metadata.update(self._extract_docx_metadata(docx_bytes))

        return metadata

    async def _extract_pdf_metadata(self, pdf_bytes: bytes) -> dict:
        """Extract metadata from PDF."""
        try:
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                metadata = {
                    "page_count": doc.page_count,
                    "has_images": any(page.get_images() for page in doc),
                    "is_encrypted": doc.is_encrypted,
                    "has_forms": any(page.get_widgets() for page in doc),
                    "metadata": doc.metadata,
                }

                # Check if PDF is scanned
                metadata["is_scanned"] = self._is_pdf_scanned(doc)

                return metadata

        except Exception as e:
            logger.error(f"Failed to extract PDF metadata: {e}")
            return {}

    def _extract_docx_metadata(self, docx_bytes: bytes) -> dict:
        """Extract metadata from DOCX."""
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            core_props = doc.core_properties

            metadata = {
                "page_count": len(doc.paragraphs),
                "has_images": False,  # DOCX images are more complex to detect
                "is_encrypted": False,
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": core_props.created or "",
                "modified": core_props.modified or "",
                "metadata": {
                    "language": core_props.language or "",
                    "category": core_props.category or "",
                    "keywords": core_props.keywords or "",
                },
            }

            return metadata

        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {e}")
            return {}


# Factory function for creating document processor
def create_document_processor() -> DocumentProcessor:
    """Create a document processor instance."""
    return DocumentProcessor()
